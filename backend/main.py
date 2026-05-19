import base64
import json
import logging
import os
import threading
import time
import uuid
import tempfile
from pathlib import Path

from faster_whisper import WhisperModel, BatchedInferencePipeline
from fastapi import FastAPI, File, UploadFile, HTTPException
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from docx import Document
import uvicorn

# ── Logging ─────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("transcriber")
logging.getLogger("faster_whisper").setLevel(logging.DEBUG)
logging.getLogger("ctranslate2").setLevel(logging.DEBUG)

# ── App ──────────────────────────────────────────────────────────────────────
app = FastAPI(title="Tamil Audio Transcriber")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Model load ───────────────────────────────────────────────────────────────
# MODAL_TASK_ID is set inside Modal containers — skip load here to avoid
# double-loading. modal_deploy.py injects the model via main.model.
model = None

if not os.environ.get("MODAL_TASK_ID"):
    log.info("=" * 60)
    log.info("Loading model: Systran/faster-whisper-large-v3")
    log.info("Device: auto  |  Compute type: int8 (CPU-optimised)")
    log.info("=" * 60)
    _t0 = time.time()
    _base = WhisperModel(
        "Systran/faster-whisper-large-v3",
        device="auto",
        compute_type="int8",        # fastest on CPU
        cpu_threads=os.cpu_count(),
        num_workers=2,
    )
    model = BatchedInferencePipeline(model=_base)
    log.info("=" * 60)
    log.info(f"Model ready — loaded in {time.time() - _t0:.1f}s")
    log.info("=" * 60)

# ── Storage ───────────────────────────────────────────────────────────────────
TEMP_DIR = Path(tempfile.gettempdir()) / "tamil_transcriber"
TEMP_DIR.mkdir(exist_ok=True)

# Jobs dir: shared Modal Volume in production, local tmp in dev
_JOBS_DIR = Path("/model-cache/jobs") if os.path.isdir("/model-cache") else TEMP_DIR / "jobs"
_JOBS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTS = {".mp3", ".wav", ".ogg", ".flac", ".m4a", ".mp4", ".webm", ".aac"}

# In-memory store — primary source for same-container requests (no NFS lag)
_mem_jobs: dict = {}


def _write_job(job_id: str, data: dict):
    _mem_jobs[job_id] = data
    try:
        p = _JOBS_DIR / f"{job_id}.json"
        p.write_text(json.dumps(data))
        # fsync so NFS flushes to storage immediately (fixes cross-container stale reads)
        with open(p, "a") as f:
            os.fsync(f.fileno())
    except Exception as e:
        log.warning(f"Job file write failed for {job_id}: {e}")


def _read_job(job_id: str) -> dict | None:
    if job_id in _mem_jobs:
        return _mem_jobs[job_id]
    p = _JOBS_DIR / f"{job_id}.json"
    if p.exists():
        try:
            return json.loads(p.read_text())
        except Exception:
            return None
    return None


# ── LLM Tamil correction ─────────────────────────────────────────────────────
def _llm_correct_tamil(text: str, emit) -> str:
    """
    Send raw Whisper output to Claude Haiku for Tamil grammar/homophone correction.
    Chunks the text to stay within output token limits.
    """
    import anthropic

    SYSTEM = (
        "You are an expert Tamil language editor. "
        "The text was auto-transcribed from Tamil audio and may contain "
        "homophone confusions, grammar errors, or misheard words. "
        "Correct them based on context. Keep everything in Tamil. "
        "Return ONLY the corrected Tamil text — every sentence, nothing omitted, no commentary."
    )

    CHUNK_CHARS = 3000
    lines = text.split("\n")
    chunks, buf, buf_len = [], [], 0
    for line in lines:
        if buf and buf_len + len(line) > CHUNK_CHARS:
            chunks.append("\n".join(buf))
            buf, buf_len = [line], len(line)
        else:
            buf.append(line)
            buf_len += len(line)
    if buf:
        chunks.append("\n".join(buf))

    total = len(chunks)
    emit(f"Correcting {total} chunk(s) with Claude Haiku (parallel)...")

    try:
        from concurrent.futures import ThreadPoolExecutor, as_completed

        client = anthropic.Anthropic()
        results = [None] * total

        def _correct(idx: int, chunk: str):
            msg = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=4096,
                system=SYSTEM,
                messages=[{"role": "user", "content": chunk}],
            )
            return idx, msg.content[0].text.strip()

        completed = 0
        with ThreadPoolExecutor(max_workers=min(total, 5)) as pool:
            futures = {pool.submit(_correct, i, c): i for i, c in enumerate(chunks)}
            for fut in as_completed(futures):
                idx, corrected = fut.result()
                results[idx] = corrected
                completed += 1
                if total > 1:
                    emit(f"  Chunk {completed}/{total} done.")

        result = "\n".join(results)
        emit(f"Claude correction complete — {len(result)} chars.")
        return result

    except Exception as exc:
        emit(f"Claude correction failed ({exc}) — using original transcription.")
        return text


# ── Transcription worker ──────────────────────────────────────────────────────
def _transcribe_job(job_id: str, audio_path: str, initial_logs: list | None = None):
    logs = list(initial_logs) if initial_logs else []

    def emit(msg: str):
        log.info(f"[{job_id}] {msg}")
        logs.append(msg)
        job = _read_job(job_id) or {}
        job["logs"] = logs
        # Strip docx_b64 from intermediate writes — keep status JSON small
        job.pop("docx_b64", None)
        _write_job(job_id, job)

    try:
        if model is None:
            raise RuntimeError("Model not loaded.")

        emit("Loading audio file...")

        segments_iter, info = model.transcribe(
            audio_path,
            language="ta",
            batch_size=16,                   # GPU batching (ignored on CPU)
            beam_size=1,                     # greedy — 3-4x faster than beam=5
            condition_on_previous_text=False, # skip context conditioning
            vad_filter=True,
            vad_parameters={"min_silence_duration_ms": 300},
        )

        emit(f"Audio: {info.duration:.2f}s  |  "
             f"Language: {info.language}  |  "
             f"Confidence: {info.language_probability:.1%}")
        emit("Decoding segments...")

        segments_list = []
        for seg in segments_iter:
            line = (f"[{seg.start:.2f}s → {seg.end:.2f}s]  "
                    f"logprob={seg.avg_logprob:.3f}  "
                    f"no_speech={seg.no_speech_prob:.3f}  "
                    f"→ {seg.text.strip()}")
            emit(line)
            segments_list.append(seg)

        full_text = "\n".join(seg.text.strip() for seg in segments_list if seg.text.strip())

        if not full_text:
            _write_job(job_id, {"status": "error", "error": "No speech detected.", "logs": logs})
            return

        emit(f"Complete — {len(segments_list)} segment(s)  |  {len(full_text)} chars")

        # LLM post-processing: fix Tamil homophones, grammar, and context
        corrected_text = _llm_correct_tamil(full_text, emit)

        # Build Word document
        doc = Document()
        doc.add_heading("Tamil Audio Transcription", level=1)

        p1 = doc.add_paragraph()
        p1.add_run("Language: ").bold = True
        p1.add_run("Tamil (தமிழ்)")

        p2 = doc.add_paragraph()
        p2.add_run("Audio Duration: ").bold = True
        p2.add_run(f"{info.duration:.1f} seconds")

        doc.add_paragraph()
        doc.add_heading("Transcription", level=2)
        para = doc.add_paragraph(corrected_text)
        para.runs[0].font.size = __import__("docx").shared.Pt(13)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        docx_b64 = base64.b64encode(buf.getvalue()).decode()
        emit("Word document ready to download.")

        # Write final state — docx_b64 stored here but NOT returned by /status
        _write_job(job_id, {
            "status": "done",
            "text": corrected_text,
            "duration": round(info.duration, 2),
            "language": "Tamil",
            "download_id": job_id,
            "docx_b64": docx_b64,   # only read by /download, never by /status
            "logs": logs,
        })

    except Exception as exc:
        log.exception(f"[{job_id}] Transcription failed: {exc}")
        _write_job(job_id, {"status": "error", "error": str(exc), "logs": logs})
    finally:
        try:
            Path(audio_path).unlink(missing_ok=True)
        except Exception:
            pass


# ── YouTube worker ───────────────────────────────────────────────────────────
def _youtube_worker(job_id: str, url: str):
    logs = []

    def emit(msg: str):
        log.info(f"[{job_id}] {msg}")
        logs.append(msg)
        job = _read_job(job_id) or {}
        job["logs"] = logs
        job.pop("docx_b64", None)
        _write_job(job_id, job)

    cookie_path = None
    audio_path  = str(TEMP_DIR / f"{job_id}.mp3")
    try:
        import yt_dlp
        import tempfile

        # Write the YOUTUBE_COOKIES secret to a temp file for yt-dlp
        cookies_env = os.environ.get("YOUTUBE_COOKIES", "").strip()
        if not cookies_env:
            raise RuntimeError(
                "No YouTube cookies found. "
                "Please update the 'youtube-cookies' Modal secret with your browser cookies."
            )
        tf = tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False)
        tf.write(cookies_env)
        tf.close()
        cookie_path = tf.name
        emit("Cookies loaded. Fetching YouTube audio...")

        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": str(TEMP_DIR / f"{job_id}.%(ext)s"),
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
                "preferredquality": "128",
            }],
            "cookiefile": cookie_path,
            "quiet": True,
            "no_warnings": True,
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            title    = info.get("title", "Unknown")
            duration = info.get("duration", 0)
            emit(f'Downloaded: "{title}"  |  {duration}s')

        emit("Starting transcription...")
        _transcribe_job(job_id, audio_path, initial_logs=logs)

    except Exception as exc:
        log.exception(f"[{job_id}] YouTube worker failed: {exc}")
        _write_job(job_id, {"status": "error", "error": str(exc), "logs": logs})
    finally:
        if cookie_path:
            try:
                os.unlink(cookie_path)
            except Exception:
                pass
        # clean up any leftover raw files yt-dlp may have left
        for p in TEMP_DIR.glob(f"{job_id}.*"):
            if str(p) != audio_path:
                p.unlink(missing_ok=True)


# ── Routes ────────────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {"status": "ok", "model": "openai/whisper-large-v2"}


@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    log.info("─" * 60)
    log.info(f"Received: '{file.filename}'  ({file.content_type})")

    ext = Path(file.filename or "").suffix.lower()
    if ext and ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail="Unsupported file type.")

    job_id = uuid.uuid4().hex
    audio_path = TEMP_DIR / f"{job_id}{ext or '.wav'}"

    contents = await file.read()
    audio_path.write_bytes(contents)
    log.info(f"[{job_id}] Saved {len(contents)/1024:.1f} KB")

    _write_job(job_id, {"status": "processing", "logs": []})

    threading.Thread(
        target=_transcribe_job,
        args=(job_id, str(audio_path)),
        daemon=True,
    ).start()

    log.info(f"[{job_id}] Transcription thread started.")
    return {"job_id": job_id, "status": "processing"}


class YoutubeReq(BaseModel):
    url: str


@app.post("/transcribe-youtube")
async def transcribe_youtube(body: YoutubeReq):
    import re
    url = body.url.strip()
    if not re.search(r"youtube\.com|youtu\.be", url):
        raise HTTPException(status_code=400, detail="Please provide a valid YouTube URL.")

    job_id = uuid.uuid4().hex
    _write_job(job_id, {"status": "processing", "logs": []})

    threading.Thread(
        target=_youtube_worker,
        args=(job_id, url),
        daemon=True,
    ).start()

    log.info(f"[{job_id}] YouTube transcription started: {url}")
    return {"job_id": job_id, "status": "processing"}


@app.get("/status/{job_id}")
def status(job_id: str):
    if not all(c in "0123456789abcdef" for c in job_id):
        raise HTTPException(status_code=400, detail="Invalid job ID.")

    job = _read_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")

    # Strip docx_b64 — keeps response small so logs render fast in UI
    return {k: v for k, v in job.items() if k != "docx_b64"}


@app.get("/download/{download_id}")
def download(download_id: str):
    if not all(c in "0123456789abcdef" for c in download_id):
        raise HTTPException(status_code=400, detail="Invalid download ID.")

    job = _read_job(download_id)
    if not job or job.get("status") != "done" or "docx_b64" not in job:
        raise HTTPException(status_code=404, detail="File not found or not ready yet.")

    log.info(f"Serving download for job: {download_id}")
    return Response(
        content=base64.b64decode(job["docx_b64"]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=tamil_transcription.docx"},
    )


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)
